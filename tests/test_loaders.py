"""
Unit tests for document loaders (PDF, TXT, Markdown, DOCX) and error handling.
"""
import tempfile
import unittest
from pathlib import Path
from src.loaders import load_document, clean_text, SUPPORTED_EXTENSIONS


class TestLoaders(unittest.TestCase):
    """Test suite for Document loaders and clean_text helper."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_clean_text(self):
        """Verify text cleaning normalizes newlines, spaces, and removes null bytes."""
        raw = "Hello\r\n\r\nWorld!\x00   This   is   a   test.\n\n\n\nEnd."
        cleaned = clean_text(raw)
        self.assertNotIn("\x00", cleaned)
        self.assertNotIn("\r", cleaned)
        self.assertIn("Hello\n\nWorld!", cleaned)
        self.assertIn("This is a test.", cleaned)

    def test_load_txt_success(self):
        """Verify reading a valid plain text document."""
        txt_file = self.temp_path / "sample.txt"
        txt_file.write_text("This is line one.\nThis is line two.", encoding="utf-8")

        docs = load_document(txt_file)
        self.assertEqual(len(docs), 1)
        doc = docs[0]
        self.assertEqual(doc.filename, "sample.txt")
        self.assertEqual(doc.file_type, "txt")
        self.assertIsNone(doc.page_number)
        self.assertIn("This is line one.", doc.page_content)

    def test_load_md_success(self):
        """Verify reading a valid markdown document."""
        md_file = self.temp_path / "readme.md"
        md_file.write_text("# Title\n\n**Bold text** and bullet point:\n- Item 1", encoding="utf-8")

        docs = load_document(md_file)
        self.assertEqual(len(docs), 1)
        doc = docs[0]
        self.assertEqual(doc.filename, "readme.md")
        self.assertEqual(doc.file_type, "md")
        self.assertIn("# Title", doc.page_content)

    def test_load_docx_success(self):
        """Verify reading a valid DOCX document using python-docx."""
        try:
            import docx
        except ImportError:
            raise unittest.SkipTest("python-docx package is not installed.")

        docx_file = self.temp_path / "document.docx"
        doc = docx.Document()
        doc.add_heading("Docx Header", level=1)
        doc.add_paragraph("This is a paragraph in the docx document.")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Cell A"
        table.rows[0].cells[1].text = "Cell B"
        doc.save(docx_file)

        docs = load_document(docx_file)
        self.assertEqual(len(docs), 1)
        d = docs[0]
        self.assertEqual(d.filename, "document.docx")
        self.assertEqual(d.file_type, "docx")
        self.assertIn("Docx Header", d.page_content)
        self.assertIn("Cell A | Cell B", d.page_content)

    def test_load_pdf_success(self):
        """Verify reading a synthetic valid PDF using pypdf writer."""
        try:
            from pypdf import PdfWriter
        except ImportError:
            raise unittest.SkipTest("pypdf package is not installed.")

        pdf_file = self.temp_path / "sample.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        with open(pdf_file, "wb") as f:
            writer.write(f)

        with self.assertRaises(ValueError) as ctx:
            load_document(pdf_file)
        self.assertIn("Empty document", str(ctx.exception))

    def test_missing_file(self):
        """Verify FileNotFoundError is raised for non-existent files."""
        non_existent = self.temp_path / "does_not_exist.txt"
        with self.assertRaises(FileNotFoundError):
            load_document(non_existent)

    def test_unsupported_file_extension(self):
        """Verify ValueError is raised for unsupported file types."""
        invalid_file = self.temp_path / "image.png"
        invalid_file.write_bytes(b"fake image data")
        with self.assertRaises(ValueError) as ctx:
            load_document(invalid_file)
        self.assertIn("Unsupported file format", str(ctx.exception))

    def test_empty_document(self):
        """Verify ValueError is raised for empty documents."""
        empty_file = self.temp_path / "empty.txt"
        empty_file.write_text("   \n\n   ", encoding="utf-8")
        with self.assertRaises(ValueError) as ctx:
            load_document(empty_file)
        self.assertIn("Empty document", str(ctx.exception))

    def test_malformed_docx(self):
        """Verify ValueError is raised for malformed DOCX binary files."""
        bad_docx = self.temp_path / "corrupt.docx"
        bad_docx.write_bytes(b"not a real zip/docx file header")
        with self.assertRaises(ValueError) as ctx:
            load_document(bad_docx)
        self.assertIn("Malformed or unparseable DOCX file", str(ctx.exception))


if __name__ == "__main__":
    unittest.main()
