"""
Modular document loaders for PDF, TXT, Markdown, and DOCX files.
Handles text extraction, structured table parsing (pdfplumber), metadata preservation, and validation.
"""
import hashlib
import re
from pathlib import Path
from typing import List, Union
from src.config import logger
from src.models import Document

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


def clean_text(text: str) -> str:
    """Normalizes whitespace and removes null characters while retaining structure."""
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\t", " ")
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def generate_doc_id(file_path: Path) -> str:
    """Generates a stable, deterministic unique document ID based on resolved path."""
    resolved_path = str(file_path.resolve()).encode("utf-8")
    return hashlib.sha256(resolved_path).hexdigest()[:16]


def format_table_as_markdown(table: List[List[Union[str, None]]]) -> str:
    """Converts a 2D list table extracted from PDF/DOCX into a Markdown table string."""
    if not table or not any(table):
        return ""
    
    clean_rows = []
    for row in table:
        if not row:
            continue
        cells = [str(c).replace("\n", " ").strip() if c is not None else "" for c in row]
        if any(cells):
            clean_rows.append(" | ".join(cells))

    if not clean_rows:
        return ""

    # Build markdown table header separator
    header_cells = clean_rows[0].split(" | ")
    sep = " | ".join(["---"] * len(header_cells))
    
    if len(clean_rows) > 1:
        return f"| {clean_rows[0]} |\n| {sep} |\n" + "\n".join(f"| {r} |" for r in clean_rows[1:])
    return f"| {clean_rows[0]} |"


def load_pdf(file_path: Path, doc_id: str) -> List[Document]:
    """Extracts text page-by-page from a PDF file using pdfplumber with pypdf fallback."""
    documents = []
    
    # Try pdfplumber first for table-aware text extraction
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages, start=1):
                text_parts = []
                
                # Extract page text
                raw_text = page.extract_text() or ""
                if raw_text:
                    text_parts.append(raw_text)

                # Extract and format tables into Markdown
                try:
                    tables = page.extract_tables()
                    for tbl in tables:
                        tbl_md = format_table_as_markdown(tbl)
                        if tbl_md:
                            text_parts.append(f"\n[Structured Table]\n{tbl_md}\n")
                except Exception:
                    pass

                full_page_text = clean_text("\n\n".join(text_parts))
                if full_page_text:
                    meta = {
                        "doc_id": doc_id,
                        "filename": file_path.name,
                        "file_type": "pdf",
                        "source_path": str(file_path.resolve()),
                        "page_number": i,
                        "total_pages": total_pages,
                    }
                    documents.append(
                        Document(
                            doc_id=doc_id,
                            filename=file_path.name,
                            file_type="pdf",
                            source_path=str(file_path.resolve()),
                            page_content=full_page_text,
                            page_number=i,
                            metadata=meta,
                        )
                    )
        if documents:
            return documents
    except Exception as e:
        logger.debug(f"pdfplumber extraction skipped for '{file_path.name}': {e}. Using pypdf fallback.")

    # Fallback to pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)

        for i, page in enumerate(reader.pages, start=1):
            try:
                raw_text = page.extract_text() or ""
            except Exception:
                raw_text = ""

            cleaned = clean_text(raw_text)
            if cleaned:
                meta = {
                    "doc_id": doc_id,
                    "filename": file_path.name,
                    "file_type": "pdf",
                    "source_path": str(file_path.resolve()),
                    "page_number": i,
                    "total_pages": total_pages,
                }
                documents.append(
                    Document(
                        doc_id=doc_id,
                        filename=file_path.name,
                        file_type="pdf",
                        source_path=str(file_path.resolve()),
                        page_content=cleaned,
                        page_number=i,
                        metadata=meta,
                    )
                )

        return documents
    except Exception as e:
        raise ValueError(f"Malformed or unparseable PDF file '{file_path.name}': {str(e)}") from e


def load_docx(file_path: Path, doc_id: str) -> List[Document]:
    """Extracts text from paragraphs and tables in a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
    except Exception as e:
        raise ValueError(f"Malformed or unparseable DOCX file '{file_path.name}': {str(e)}") from e

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
        tbl_md = format_table_as_markdown(table_rows)
        if tbl_md:
            text_parts.append(f"\n[Structured Table]\n{tbl_md}\n")

    full_text = clean_text("\n\n".join(text_parts))
    if not full_text:
        return []

    meta = {
        "doc_id": doc_id,
        "filename": file_path.name,
        "file_type": "docx",
        "source_path": str(file_path.resolve()),
        "page_number": None,
    }
    return [
        Document(
            doc_id=doc_id,
            filename=file_path.name,
            file_type="docx",
            source_path=str(file_path.resolve()),
            page_content=full_text,
            page_number=None,
            metadata=meta,
        )
    ]


def load_txt(file_path: Path, doc_id: str, file_type: str = "txt") -> List[Document]:
    """Extracts text from a plain TXT or Markdown file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()
    except Exception as e:
        raise ValueError(f"Failed to read file '{file_path.name}': {str(e)}") from e

    cleaned = clean_text(raw_text)
    if not cleaned:
        return []

    meta = {
        "doc_id": doc_id,
        "filename": file_path.name,
        "file_type": file_type,
        "source_path": str(file_path.resolve()),
        "page_number": None,
    }
    return [
        Document(
            doc_id=doc_id,
            filename=file_path.name,
            file_type=file_type,
            source_path=str(file_path.resolve()),
            page_content=cleaned,
            page_number=None,
            metadata=meta,
        )
    ]


def load_document(file_path: Union[str, Path]) -> List[Document]:
    """
    Factory loader function for enterprise documents.
    Validates file existence, format support, and extracted content.
    """
    path = Path(file_path).resolve()

    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"Source file does not exist: {path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format '{ext}' for file '{path.name}'. Supported formats: {sorted(list(SUPPORTED_EXTENSIONS))}"
        )

    doc_id = generate_doc_id(path)
    logger.info(f"Loading document '{path.name}' (type: {ext}, doc_id: {doc_id})")

    if ext == ".pdf":
        documents = load_pdf(path, doc_id)
    elif ext == ".docx":
        documents = load_docx(path, doc_id)
    elif ext == ".md":
        documents = load_txt(path, doc_id, file_type="md")
    else:  # .txt
        documents = load_txt(path, doc_id, file_type="txt")

    if not documents:
        raise ValueError(f"Empty document or no readable text content in '{path.name}'.")

    return documents
